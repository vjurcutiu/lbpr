from __future__ import annotations

from dataclasses import dataclass, field
from io import BytesIO
import re
from typing import Iterable, Literal
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_ORIENTATION, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    ListFlowable,
    ListItem,
    NextPageTemplate,
    PageBreak,
    PageTemplate,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

MarkdownExportFormat = Literal["markdown", "txt", "docx", "pdf"]
MarkdownBlockKind = Literal["heading", "paragraph", "bullet", "numbered", "table"]


_INTERNAL_SECTION_HEADINGS = {
    "workflow notes",
    "source notes",
    "internal notes",
    "app notes",
}


LEGAL_EXPORT_FIELD_LABELS: dict[str, str] = {
    "business_impact": "Business Impact",
    "clause_family": "Clause Type",
    "clause_type": "Clause Type",
    "current_position": "Current Position",
    "fallback_position": "Fallback Position",
    "follow_up": "Follow-Up",
    "recommended_change": "Recommended Change",
    "recommended_position": "Recommended Position",
    "requires_human_review": "Requires Human Review",
    "responsible_party": "Responsible Party",
    "risk_note": "Risk Note",
    "source_basis": "Source Basis",
    "trigger_or_deadline": "Trigger / Deadline",
}

LEGAL_CLAUSE_TYPE_LABELS: dict[str, str] = {
    "assignment": "Assignment",
    "audit": "Audit",
    "change_control": "Change Control",
    "confidentiality": "Confidentiality",
    "data_protection": "Data Protection",
    "dispute_resolution": "Dispute Resolution",
    "exclusivity": "Exclusivity",
    "force_majeure": "Force Majeure",
    "governing_law": "Governing Law",
    "indemnity": "Indemnity",
    "insurance": "Insurance",
    "ip_ownership": "IP Ownership",
    "intellectual_property": "Intellectual Property",
    "limitation_of_liability": "Limitation of Liability",
    "non_compete": "Non-Compete",
    "non_solicit": "Non-Solicit",
    "notices": "Notices",
    "payment": "Payment",
    "renewal": "Renewal",
    "residuals": "Residuals",
    "service_levels": "Service Levels",
    "termination": "Termination",
    "warranties": "Warranties",
}

LEGAL_EXPORT_TOKEN_LABELS: dict[str, str] = {**LEGAL_EXPORT_FIELD_LABELS, **LEGAL_CLAUSE_TYPE_LABELS}
LEGAL_EXPORT_FIELD_LABEL_KEYS: dict[str, str] = {
    re.sub(r"[^a-z0-9]+", "_", key.strip().lower()).strip("_"): label
    for key, label in LEGAL_EXPORT_FIELD_LABELS.items()
}

_INTERNAL_CLAUSE_TOKEN_RE = re.compile(
    r"\b("
    + "|".join(
        re.escape(key)
        for key in sorted(LEGAL_EXPORT_TOKEN_LABELS, key=len, reverse=True)
        if "_" in key
    )
    + r")\b",
    re.IGNORECASE,
)


def _normalize_export_label_key(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", str(value or "").strip().lower()).strip("_")


def _normalize_clause_type_key(value: str) -> str:
    return _normalize_export_label_key(value)


def _format_export_field_label(value: str) -> str:
    raw = str(value or "").strip()
    return LEGAL_EXPORT_FIELD_LABEL_KEYS.get(_normalize_export_label_key(raw), raw)


def _format_clause_type_label(value: str) -> str:
    raw = str(value or "").strip()
    key = _normalize_clause_type_key(raw)
    return LEGAL_CLAUSE_TYPE_LABELS.get(key, raw)


def _format_internal_clause_tokens(text: str) -> str:
    return _INTERNAL_CLAUSE_TOKEN_RE.sub(
        lambda match: LEGAL_EXPORT_TOKEN_LABELS.get(_normalize_clause_type_key(match.group(0)), match.group(0)),
        str(text or ""),
    )


def _format_clause_table_cell(text: str) -> str:
    raw = str(text or "").strip()
    if not raw:
        return raw

    field_label = _format_export_field_label(raw)
    if field_label != raw:
        return field_label

    exact = _format_clause_type_label(raw)
    if exact != raw:
        return exact

    parts = [part.strip() for part in re.split(r"\s*[,;]\s*", raw) if part.strip()]
    if len(parts) > 1:
        normalized = [_normalize_clause_type_key(part) for part in parts]
        if all(part in LEGAL_CLAUSE_TYPE_LABELS for part in normalized):
            return ", ".join(LEGAL_CLAUSE_TYPE_LABELS[part] for part in normalized)

    return _format_internal_clause_tokens(raw)


def _format_export_clause_labels(markdown: str) -> str:
    formatted_lines: list[str] = []
    for line in str(markdown or "").split("\n"):
        if _is_table_line(line):
            cells = _split_table_cells(line)
            if cells and _is_table_separator_row(cells):
                formatted_lines.append(line)
            elif cells:
                formatted_lines.append("| " + " | ".join(cells) + " |")
            else:
                formatted_lines.append(_format_internal_clause_tokens(line))
        else:
            formatted_lines.append(_format_internal_clause_tokens(line))
    return "\n".join(formatted_lines)


def sanitize_export_markdown(markdown: str) -> str:
    text = str(markdown or "").replace("\r\n", "\n").replace("\r", "\n")
    lines = text.split("\n")
    cleaned: list[str] = []
    skipping = False
    skip_level = 0

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()
        heading_match = _HEADING_RE.match(stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = _strip_inline_markdown(heading_match.group(2)).strip().lower()
            if skipping and level <= skip_level:
                skipping = False
                skip_level = 0
            if heading_text in _INTERNAL_SECTION_HEADINGS:
                skipping = True
                skip_level = level
                continue
        if skipping:
            continue
        cleaned.append(line)

    sanitized = "\n".join(cleaned)
    sanitized = re.sub(r"\n{3,}", "\n\n", sanitized)
    return sanitized.strip()


@dataclass(slots=True)
class ExportedArtifact:
    content: bytes
    file_name: str
    content_type: str
    format: MarkdownExportFormat


@dataclass(slots=True)
class MarkdownBlock:
    kind: MarkdownBlockKind
    text: str = ""
    level: int = 0
    headers: list[str] = field(default_factory=list)
    rows: list[list[str]] = field(default_factory=list)


def export_artifact(*, title: str, markdown: str, file_stem: str, target_format: MarkdownExportFormat) -> ExportedArtifact:
    normalized = _normalize_markdown(_format_export_clause_labels(sanitize_export_markdown(markdown)))
    if target_format == "markdown":
        return ExportedArtifact(
            content=normalized.encode("utf-8"),
            file_name=f"{file_stem}.md",
            content_type="text/markdown; charset=utf-8",
            format=target_format,
        )

    blocks = _parse_markdown_blocks(normalized)
    if target_format == "txt":
        plain = _render_plain_text(blocks, fallback_title=title)
        return ExportedArtifact(
            content=plain.encode("utf-8"),
            file_name=f"{file_stem}.txt",
            content_type="text/plain; charset=utf-8",
            format=target_format,
        )
    if target_format == "docx":
        payload = _render_docx(title=title, blocks=blocks)
        return ExportedArtifact(
            content=payload,
            file_name=f"{file_stem}.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            format=target_format,
        )
    if target_format == "pdf":
        payload = _render_pdf(title=title, blocks=blocks)
        return ExportedArtifact(
            content=payload,
            file_name=f"{file_stem}.pdf",
            content_type="application/pdf",
            format=target_format,
        )
    raise ValueError(f"Unsupported artifact export format: {target_format}")


_INLINE_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
_INLINE_IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")


def _normalize_markdown(markdown: str) -> str:
    text = str(markdown or "").replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\u00a0", " ")
    return text.strip()


def _strip_inline_markdown(text: str) -> str:
    clean = str(text or "")
    clean = _INLINE_IMAGE_RE.sub(lambda m: (m.group(1) or m.group(2) or "").strip(), clean)
    clean = _INLINE_LINK_RE.sub(lambda m: f"{m.group(1).strip()} ({m.group(2).strip()})", clean)
    clean = re.sub(r"`([^`]+)`", r"\1", clean)
    clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", clean)
    clean = re.sub(r"__([^_]+)__", r"\1", clean)
    clean = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", clean)
    clean = re.sub(r"(?<!_)_([^_]+)_(?!_)", r"\1", clean)
    clean = re.sub(r"~~([^~]+)~~", r"\1", clean)
    clean = re.sub(r"^>+\s?", "", clean)
    clean = re.sub(r"\s+", " ", clean)
    return clean.strip()


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET_RE = re.compile(r"^[-*+]\s+(.*)$")
_NUMBERED_RE = re.compile(r"^\d+[.)]\s+(.*)$")
_TABLE_SEPARATOR_CELL_RE = re.compile(r"^:?-{3,}:?$")


def _is_table_line(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if stripped.startswith("|") and stripped.count("|") >= 2:
        return True
    return stripped.count("|") >= 2 and " | " in stripped


def _is_table_separator_row(cells: list[str]) -> bool:
    meaningful = [cell.strip() for cell in cells if cell.strip()]
    if not meaningful:
        return False
    return all(bool(_TABLE_SEPARATOR_CELL_RE.match(cell)) for cell in meaningful)


def _split_table_cells(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    cells = re.split(r"(?<!\\)\|", stripped)
    return [_format_clause_table_cell(_strip_inline_markdown(cell.replace(r"\|", "|").strip())) for cell in cells]


def _parse_table_lines(lines: list[str]) -> MarkdownBlock | None:
    parsed_rows = [_split_table_cells(line) for line in lines if _is_table_line(line)]
    parsed_rows = [row for row in parsed_rows if len(row) >= 2 and any(cell for cell in row)]
    if not parsed_rows:
        return None

    header_index = 0
    header = parsed_rows[header_index]
    body_start = header_index + 1
    if len(parsed_rows) > 1 and _is_table_separator_row(parsed_rows[1]):
        body_start = 2
    elif _is_table_separator_row(header):
        if len(parsed_rows) < 2:
            return None
        header = parsed_rows[1]
        body_start = 2

    body_rows = [row for row in parsed_rows[body_start:] if not _is_table_separator_row(row)]
    if not body_rows and len(parsed_rows) <= 1:
        return None

    column_count = max(len(header), *(len(row) for row in body_rows)) if body_rows else len(header)
    if column_count < 2:
        return None

    normalized_header = _normalize_table_row(header, column_count)
    normalized_rows = [_normalize_table_row(row, column_count) for row in body_rows]
    return MarkdownBlock(kind="table", headers=normalized_header, rows=normalized_rows)


def _normalize_table_row(row: list[str], column_count: int) -> list[str]:
    values = list(row[:column_count])
    if len(values) < column_count:
        values.extend([""] * (column_count - len(values)))
    return values


def _parse_markdown_blocks(markdown: str) -> list[MarkdownBlock]:
    blocks: list[MarkdownBlock] = []
    paragraph_lines: list[str] = []
    in_code_fence = False
    lines = markdown.split("\n")
    index = 0

    def flush_paragraph() -> None:
        if not paragraph_lines:
            return
        text = _strip_inline_markdown(" ".join(paragraph_lines))
        paragraph_lines.clear()
        if text:
            blocks.append(MarkdownBlock(kind="paragraph", text=text))

    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code_fence = not in_code_fence
            flush_paragraph()
            index += 1
            continue
        if in_code_fence:
            if stripped:
                paragraph_lines.append(stripped)
            index += 1
            continue
        if not stripped:
            flush_paragraph()
            index += 1
            continue

        if _is_table_line(stripped):
            table_lines: list[str] = []
            while index < len(lines) and _is_table_line(lines[index].strip()):
                table_lines.append(lines[index].rstrip())
                index += 1
            table = _parse_table_lines(table_lines)
            if table:
                flush_paragraph()
                blocks.append(table)
            else:
                paragraph_lines.extend(line.strip() for line in table_lines if line.strip())
            continue

        heading_match = _HEADING_RE.match(stripped)
        if heading_match:
            flush_paragraph()
            text = _strip_inline_markdown(heading_match.group(2))
            if text:
                blocks.append(MarkdownBlock(kind="heading", text=text, level=len(heading_match.group(1))))
            index += 1
            continue

        bullet_match = _BULLET_RE.match(stripped)
        if bullet_match:
            flush_paragraph()
            text = _strip_inline_markdown(bullet_match.group(1))
            if text:
                blocks.append(MarkdownBlock(kind="bullet", text=text))
            index += 1
            continue

        numbered_match = _NUMBERED_RE.match(stripped)
        if numbered_match:
            flush_paragraph()
            text = _strip_inline_markdown(numbered_match.group(1))
            if text:
                blocks.append(MarkdownBlock(kind="numbered", text=text))
            index += 1
            continue

        paragraph_lines.append(stripped)
        index += 1

    flush_paragraph()
    return blocks


def _render_plain_text(blocks: list[MarkdownBlock], *, fallback_title: str) -> str:
    if not blocks and fallback_title:
        return fallback_title.strip()

    lines: list[str] = []
    for idx, block in enumerate(blocks):
        if block.kind == "heading":
            text = block.text.strip()
            lines.append(text)
            if block.level <= 2:
                underline = "=" if block.level == 1 else "-"
                lines.append(underline * len(text))
        elif block.kind == "bullet":
            lines.append(f"- {block.text}")
        elif block.kind == "numbered":
            lines.append(f"1. {block.text}")
        elif block.kind == "table":
            lines.extend(_render_plain_text_table(block))
        else:
            lines.append(block.text)
        if idx < len(blocks) - 1:
            lines.append("")
    return "\n".join(lines).strip()


def _render_plain_text_table(block: MarkdownBlock) -> list[str]:
    rows = [block.headers, *block.rows]
    if not rows:
        return []
    column_count = max(len(row) for row in rows)
    normalized_rows = [_normalize_table_row(row, column_count) for row in rows]
    widths = [
        max(len(row[col]) for row in normalized_rows)
        for col in range(column_count)
    ]

    def render_row(row: list[str]) -> str:
        return "| " + " | ".join(row[col].ljust(widths[col]) for col in range(column_count)) + " |"

    separator = "| " + " | ".join("-" * max(width, 3) for width in widths) + " |"
    rendered = [render_row(normalized_rows[0]), separator]
    rendered.extend(render_row(row) for row in normalized_rows[1:])
    return rendered


def _apply_doc_styles(document: Document) -> None:
    _configure_section(document.sections[0], landscape_mode=False)
    document.sections[0].start_type = WD_SECTION.NEW_PAGE

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10.5)

    for style_name, size, bold in [
        ("Title", 18, True),
        ("Heading 1", 14, True),
        ("Heading 2", 12, True),
        ("Heading 3", 11, True),
    ]:
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = bold


def _configure_section(section, *, landscape_mode: bool) -> None:
    section.orientation = WD_ORIENTATION.LANDSCAPE if landscape_mode else WD_ORIENTATION.PORTRAIT
    if landscape_mode:
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        margin = Inches(0.45)
    else:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        margin = Inches(0.75)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin


def _has_table_blocks(blocks: Iterable[MarkdownBlock]) -> bool:
    return any(block.kind == "table" for block in blocks)


def _block_needs_landscape(blocks: list[MarkdownBlock], index: int) -> bool:
    block = blocks[index]
    if block.kind == "table":
        return True
    return block.kind == "heading" and index + 1 < len(blocks) and blocks[index + 1].kind == "table"


def _switch_docx_orientation(document: Document, *, landscape_mode: bool, current_landscape: bool, has_content: bool) -> bool:
    if current_landscape == landscape_mode:
        return current_landscape
    if has_content:
        section = document.add_section(WD_SECTION.NEW_PAGE)
    else:
        section = document.sections[-1]
    _configure_section(section, landscape_mode=landscape_mode)
    return landscape_mode


def _render_docx(*, title: str, blocks: list[MarkdownBlock]) -> bytes:
    document = Document()
    _apply_doc_styles(document)
    document.core_properties.title = title

    if not blocks:
        para = document.add_paragraph(title.strip() or "Workflow output")
        para.style = document.styles["Title"]
    else:
        title_added = False
        current_landscape = False
        has_content = False
        for index, block in enumerate(blocks):
            needs_landscape = _block_needs_landscape(blocks, index)
            current_landscape = _switch_docx_orientation(
                document,
                landscape_mode=needs_landscape,
                current_landscape=current_landscape,
                has_content=has_content,
            )
            if block.kind == "heading":
                level = max(1, min(block.level, 3))
                para = document.add_paragraph(block.text)
                para.style = document.styles[f"Heading {level}"]
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                title_added = True if level == 1 else title_added
                has_content = True
            elif block.kind == "bullet":
                para = document.add_paragraph(block.text, style="List Bullet")
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                has_content = True
            elif block.kind == "numbered":
                para = document.add_paragraph(block.text, style="List Number")
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                has_content = True
            elif block.kind == "table":
                _add_docx_table(document, block)
                has_content = True
            else:
                para = document.add_paragraph(block.text)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                has_content = True
        if not title_added and title.strip():
            first = document.paragraphs[0]
            first.insert_paragraph_before(title.strip(), style="Title")

    handle = BytesIO()
    document.save(handle)
    return handle.getvalue()


def _docx_available_width_inches(document: Document) -> float:
    section = document.sections[-1]
    return float(section.page_width - section.left_margin - section.right_margin) / 914400.0


def _docx_font_size_for_columns(column_count: int) -> float:
    if column_count <= 4:
        return 9.0
    if column_count <= 6:
        return 8.3
    return 7.4


def _add_docx_table(document: Document, block: MarkdownBlock) -> None:
    headers = block.headers or []
    rows = block.rows or []
    column_count = max(len(headers), *(len(row) for row in rows)) if rows else len(headers)
    if column_count <= 0:
        return

    headers = _normalize_table_row(headers, column_count)
    rows = [_normalize_table_row(row, column_count) for row in rows]
    table = document.add_table(rows=1 + len(rows), cols=column_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    available_width = _docx_available_width_inches(document)
    col_width = max(0.65, available_width / column_count)
    font_size = _docx_font_size_for_columns(column_count)

    for col_index, header in enumerate(headers):
        cell = table.cell(0, col_index)
        _format_docx_cell(cell, header, bold=True, font_size=font_size, width_inches=col_width, shaded=True)
    _repeat_docx_table_header(table.rows[0])
    _prevent_docx_row_split(table.rows[0])

    for row_index, row in enumerate(rows, start=1):
        _prevent_docx_row_split(table.rows[row_index])
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            _format_docx_cell(cell, value, bold=False, font_size=font_size, width_inches=col_width, shaded=False)

    document.add_paragraph()


def _format_docx_cell(cell, text: str, *, bold: bool, font_size: float, width_inches: float, shaded: bool) -> None:
    cell.width = Inches(width_inches)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    _set_docx_cell_margins(cell, top=80, start=80, bottom=80, end=80)
    if shaded:
        _shade_docx_cell(cell, "EDEDED")
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text or ""))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(font_size)


def _shade_docx_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _set_docx_cell_margins(cell, *, top: int, start: int, bottom: int, end: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _repeat_docx_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _prevent_docx_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is not None:
        return
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _escape_pdf_text(text: str) -> str:
    return escape(text).replace("\n", "<br/>")


def _build_pdf_styles() -> tuple[ParagraphStyle, dict[int, ParagraphStyle], ParagraphStyle, ParagraphStyle]:
    styles = getSampleStyleSheet()
    normal = ParagraphStyle(
        "WorkflowNormal",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10,
        leading=14,
        spaceAfter=8,
        alignment=TA_LEFT,
    )
    headings = {
        1: ParagraphStyle("WorkflowH1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=16, leading=20, textColor=colors.black, spaceAfter=10),
        2: ParagraphStyle("WorkflowH2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=13, leading=17, textColor=colors.black, spaceAfter=8),
        3: ParagraphStyle("WorkflowH3", parent=styles["Heading3"], fontName="Helvetica-Bold", fontSize=11, leading=15, textColor=colors.black, spaceAfter=6),
    }
    table_cell = ParagraphStyle(
        "WorkflowTableCell",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=7.0,
        leading=8.6,
        alignment=TA_LEFT,
        splitLongWords=0,
    )
    table_header = ParagraphStyle(
        "WorkflowTableHeader",
        parent=table_cell,
        fontName="Helvetica-Bold",
        fontSize=7.1,
        leading=8.7,
        textColor=colors.black,
    )
    return normal, headings, table_cell, table_header


def _build_pdf_story(blocks: list[MarkdownBlock]) -> tuple[list[object], str]:
    normal, headings, table_cell, table_header = _build_pdf_styles()
    story: list[object] = []
    pending_bullets: list[tuple[str, str]] = []
    page_mode = "Portrait"
    initial_page_mode = "Portrait"
    initial_page_mode_set = False

    def switch_page_mode(target: str) -> None:
        nonlocal page_mode, initial_page_mode, initial_page_mode_set
        if not initial_page_mode_set:
            initial_page_mode = target
            page_mode = target
            initial_page_mode_set = True
            return
        if page_mode == target:
            return
        story.append(NextPageTemplate(target))
        story.append(PageBreak())
        page_mode = target

    def flush_list() -> None:
        nonlocal pending_bullets
        if not pending_bullets:
            return
        ordered = pending_bullets[0][0] == "numbered"
        list_items = [
            ListItem(Paragraph(_escape_pdf_text(text), normal), leftIndent=12)
            for _, text in pending_bullets
        ]
        story.append(ListFlowable(list_items, bulletType="1" if ordered else "bullet", start="1", leftIndent=16))
        story.append(Spacer(1, 0.08 * inch))
        pending_bullets = []

    for index, block in enumerate(blocks):
        target_mode = "Landscape" if _block_needs_landscape(blocks, index) else "Portrait"
        if block.kind not in {"bullet", "numbered"}:
            flush_list()
            switch_page_mode(target_mode)
        if block.kind in {"bullet", "numbered"}:
            kind = block.kind
            if pending_bullets and pending_bullets[0][0] != kind:
                flush_list()
            switch_page_mode(target_mode)
            pending_bullets.append((kind, block.text))
            continue

        if block.kind == "heading":
            style = headings.get(min(max(block.level, 1), 3), headings[3])
            story.append(Paragraph(_escape_pdf_text(block.text), style))
            story.append(Spacer(1, 0.04 * inch))
        elif block.kind == "table":
            table = _build_pdf_table(block, table_cell=table_cell, table_header=table_header)
            if table is not None:
                story.append(table)
                story.append(Spacer(1, 0.18 * inch))
        else:
            story.append(Paragraph(_escape_pdf_text(block.text), normal))
    flush_list()
    return story, initial_page_mode


def _build_pdf_table(block: MarkdownBlock, *, table_cell: ParagraphStyle, table_header: ParagraphStyle) -> Table | None:
    headers = block.headers or []
    rows = block.rows or []
    column_count = max(len(headers), *(len(row) for row in rows)) if rows else len(headers)
    if column_count <= 0:
        return None

    headers = _normalize_table_row(headers, column_count)
    rows = [_normalize_table_row(row, column_count) for row in rows]
    data = [
        [Paragraph(_escape_pdf_text(cell), table_header) for cell in headers],
        *[[Paragraph(_escape_pdf_text(cell), table_cell) for cell in row] for row in rows],
    ]

    available_width = landscape(A4)[0] - (0.5 * inch * 2)
    col_widths = _calculate_pdf_column_widths([headers, *rows], available_width)
    table = Table(data, colWidths=col_widths, repeatRows=1, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EDEDED")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#9CA3AF")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 3.5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3.5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def _table_header_key(header: str) -> str:
    normalized = _normalize_export_label_key(_format_export_field_label(header))
    if normalized == "clause_family":
        return "clause_type"
    return normalized


def _pdf_min_column_width(header: str) -> float:
    key = _table_header_key(header)
    width_inches_by_key = {
        "issue": 1.22,
        "severity": 0.58,
        "clause_type": 0.92,
        "business_impact": 1.25,
        "source_basis": 1.28,
        "recommended_fix": 1.28,
        "recommended_change": 1.28,
        "recommended_position": 1.28,
        "fallback": 1.12,
        "fallback_position": 1.12,
        "approval": 0.9,
        "approver": 0.9,
        "owner": 0.72,
        "status": 0.72,
        "current_position": 1.22,
        "responsible_party": 0.95,
        "trigger_deadline": 1.05,
        "trigger_or_deadline": 1.05,
        "follow_up": 1.0,
        "risk_note": 1.0,
    }
    return width_inches_by_key.get(key, 0.72) * inch


def _calculate_pdf_column_widths(rows: list[list[str]], available_width: float) -> list[float]:
    if not rows:
        return []
    column_count = max(len(row) for row in rows)
    normalized_rows = [_normalize_table_row(row, column_count) for row in rows]
    headers = normalized_rows[0] if normalized_rows else [""] * column_count

    min_widths = [_pdf_min_column_width(headers[col]) for col in range(column_count)]
    min_total = sum(min_widths) or 1.0
    if min_total >= available_width:
        scale = available_width / min_total
        return [width * scale for width in min_widths]

    weights: list[float] = []
    for col in range(column_count):
        values = [row[col] for row in normalized_rows]
        max_len = max(len(value) for value in values)
        avg_len = sum(len(value) for value in values) / max(len(values), 1)
        header_key = _table_header_key(headers[col])
        semantic_boost = 1.15 if header_key in {
            "issue",
            "business_impact",
            "source_basis",
            "recommended_fix",
            "recommended_change",
            "recommended_position",
            "fallback",
            "fallback_position",
            "current_position",
        } else 0.85
        weights.append(max(4.0, min(36.0, (max_len * 0.6 + avg_len * 0.4) * semantic_boost)))

    remaining_width = available_width - min_total
    total_weight = sum(weights) or 1.0
    widths = [
        min_widths[col] + remaining_width * (weights[col] / total_weight)
        for col in range(column_count)
    ]
    return widths


def _render_pdf(*, title: str, blocks: list[MarkdownBlock]) -> bytes:
    handle = BytesIO()
    story, initial_page_mode = _build_pdf_story(blocks)
    if not story:
        styles = getSampleStyleSheet()
        story = [Paragraph(_escape_pdf_text(title or "Workflow output"), styles["Title"])]
        initial_page_mode = "Portrait"

    if not _has_table_blocks(blocks):
        doc = SimpleDocTemplate(
            handle,
            pagesize=A4,
            title=title,
            leftMargin=0.75 * inch,
            rightMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
        )
        doc.build(story)
        return handle.getvalue()

    portrait_frame = Frame(0.75 * inch, 0.75 * inch, A4[0] - 1.5 * inch, A4[1] - 1.5 * inch, id="portrait-frame")
    landscape_size = landscape(A4)
    landscape_frame = Frame(0.5 * inch, 0.5 * inch, landscape_size[0] - 1.0 * inch, landscape_size[1] - 1.0 * inch, id="landscape-frame")
    portrait_template = PageTemplate(id="Portrait", frames=[portrait_frame], pagesize=A4)
    landscape_template = PageTemplate(id="Landscape", frames=[landscape_frame], pagesize=landscape_size)
    page_templates = [landscape_template, portrait_template] if initial_page_mode == "Landscape" else [portrait_template, landscape_template]

    doc = BaseDocTemplate(
        handle,
        pageTemplates=page_templates,
        title=title,
    )
    doc.build(story)
    return handle.getvalue()
